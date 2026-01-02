"""Report template engine for investment research reports."""

import asyncio
import logging
import json
import os
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, Template
import markdown
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class OutputFormat(str, Enum):
    """Supported output formats."""
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    JSON = "json"


class TemplateType(str, Enum):
    """Template types for different report styles."""
    COMPREHENSIVE = "comprehensive"
    SUMMARY = "summary"
    EXECUTIVE = "executive"
    TECHNICAL = "technical"
    QUICK = "quick"
    CUSTOM = "custom"


@dataclass
class ReportSection:
    """Represents a section in the report."""
    section_id: str
    title: str
    content: str = ""
    subsections: List['ReportSection'] = field(default_factory=list)
    order: int = 0
    include_in_toc: bool = True
    page_break_before: bool = False
    
    # Formatting options
    heading_level: int = 1
    style: Optional[str] = None
    
    # Content metadata
    word_count: int = 0
    data_sources: List[str] = field(default_factory=list)
    confidence_score: float = 0.0
    
    def calculate_word_count(self):
        """Calculate word count for the section."""
        self.word_count = len(self.content.split()) if self.content else 0
        for subsection in self.subsections:
            subsection.calculate_word_count()
            self.word_count += subsection.word_count


@dataclass
class ReportTemplate:
    """Report template definition."""
    template_id: str
    name: str
    description: str
    template_type: TemplateType
    
    # Template structure
    sections: List[str] = field(default_factory=list)  # Section IDs in order
    required_sections: List[str] = field(default_factory=list)
    optional_sections: List[str] = field(default_factory=list)
    
    # Template files
    html_template: Optional[str] = None
    css_file: Optional[str] = None
    docx_template: Optional[str] = None
    
    # Configuration
    include_toc: bool = True
    include_executive_summary: bool = True
    include_appendices: bool = False
    max_pages: Optional[int] = None
    
    # Metadata
    created_at: datetime = field(default_factory=datetime.utcnow)
    version: str = "1.0"
    author: str = "system"


@dataclass
class ReportData:
    """Data structure for report generation."""
    # Basic information
    title: str
    subtitle: Optional[str] = None
    company_name: Optional[str] = None
    ticker_symbol: Optional[str] = None
    
    # Report metadata
    report_type: str = "Investment Research"
    generated_at: datetime = field(default_factory=datetime.utcnow)
    generated_by: str = "AI Research System"
    report_id: Optional[str] = None
    
    # Content sections
    sections: Dict[str, ReportSection] = field(default_factory=dict)
    
    # Analysis results
    executive_summary: str = ""
    key_findings: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    risk_factors: List[str] = field(default_factory=list)
    
    # Data and sources
    data_sources: List[str] = field(default_factory=list)
    analysis_date: datetime = field(default_factory=datetime.utcnow)
    data_quality_score: float = 0.0
    
    # Financial data
    financial_metrics: Dict[str, Any] = field(default_factory=dict)
    market_data: Dict[str, Any] = field(default_factory=dict)
    
    # Appendices
    appendices: Dict[str, Any] = field(default_factory=dict)
    
    def add_section(self, section: ReportSection):
        """Add a section to the report."""
        self.sections[section.section_id] = section
    
    def get_section(self, section_id: str) -> Optional[ReportSection]:
        """Get a section by ID."""
        return self.sections.get(section_id)
    
    def calculate_total_word_count(self) -> int:
        """Calculate total word count for the report."""
        total = len(self.executive_summary.split()) if self.executive_summary else 0
        for section in self.sections.values():
            section.calculate_word_count()
            total += section.word_count
        return total


class ReportTemplateEngine:
    """Engine for generating reports from templates and data."""
    
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = Path(templates_dir)
        self.templates: Dict[str, ReportTemplate] = {}
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
        # Create templates directory if it doesn't exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize default templates
        self._create_default_templates()
    
    def _create_default_templates(self):
        """Create default report templates."""
        # Comprehensive template
        comprehensive = ReportTemplate(
            template_id="comprehensive",
            name="Comprehensive Investment Research Report",
            description="Full-featured investment research report with all sections",
            template_type=TemplateType.COMPREHENSIVE,
            sections=[
                "executive_summary", "company_overview", "industry_analysis",
                "financial_analysis", "market_analysis", "risk_analysis",
                "valuation", "recommendations", "appendices"
            ],
            required_sections=[
                "executive_summary", "company_overview", "financial_analysis",
                "recommendations"
            ],
            include_toc=True,
            include_executive_summary=True,
            include_appendices=True
        )
        self.templates[comprehensive.template_id] = comprehensive
        
        # Summary template
        summary = ReportTemplate(
            template_id="summary",
            name="Investment Summary Report",
            description="Concise summary report for quick decision making",
            template_type=TemplateType.SUMMARY,
            sections=[
                "executive_summary", "key_metrics", "recommendations", "risks"
            ],
            required_sections=["executive_summary", "recommendations"],
            include_toc=False,
            include_executive_summary=True,
            include_appendices=False,
            max_pages=5
        )
        self.templates[summary.template_id] = summary
        
        # Executive template
        executive = ReportTemplate(
            template_id="executive",
            name="Executive Brief",
            description="High-level executive summary for senior management",
            template_type=TemplateType.EXECUTIVE,
            sections=["executive_summary", "key_findings", "recommendations"],
            required_sections=["executive_summary", "recommendations"],
            include_toc=False,
            include_executive_summary=True,
            include_appendices=False,
            max_pages=3
        )
        self.templates[executive.template_id] = executive
        
        # Quick template
        quick = ReportTemplate(
            template_id="quick",
            name="Quick Analysis Report",
            description="Fast analysis report for urgent requests",
            template_type=TemplateType.QUICK,
            sections=["summary", "key_points", "recommendation"],
            required_sections=["summary", "recommendation"],
            include_toc=False,
            include_executive_summary=False,
            include_appendices=False,
            max_pages=2
        )
        self.templates[quick.template_id] = quick
    
    def register_template(self, template: ReportTemplate):
        """Register a new report template."""
        self.templates[template.template_id] = template
        logger.info(f"Registered template: {template.name} ({template.template_id})")
    
    def get_template(self, template_id: str) -> Optional[ReportTemplate]:
        """Get template by ID."""
        return self.templates.get(template_id)
    
    def list_templates(self) -> List[ReportTemplate]:
        """List all available templates."""
        return list(self.templates.values())
    
    async def generate_report(
        self,
        template_id: str,
        report_data: ReportData,
        output_format: OutputFormat = OutputFormat.HTML,
        output_path: Optional[str] = None
    ) -> str:
        """Generate a report using the specified template and data."""
        template = self.get_template(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        logger.info(f"Generating {output_format.value} report using template: {template.name}")
        
        # Validate required sections
        missing_sections = []
        for section_id in template.required_sections:
            if section_id not in report_data.sections and not hasattr(report_data, section_id):
                missing_sections.append(section_id)
        
        if missing_sections:
            logger.warning(f"Missing required sections: {missing_sections}")
        
        # Generate based on output format
        if output_format == OutputFormat.HTML:
            return await self._generate_html_report(template, report_data, output_path)
        elif output_format == OutputFormat.PDF:
            return await self._generate_pdf_report(template, report_data, output_path)
        elif output_format == OutputFormat.DOCX:
            return await self._generate_docx_report(template, report_data, output_path)
        elif output_format == OutputFormat.MARKDOWN:
            return await self._generate_markdown_report(template, report_data, output_path)
        elif output_format == OutputFormat.JSON:
            return await self._generate_json_report(template, report_data, output_path)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    async def _generate_html_report(
        self,
        template: ReportTemplate,
        report_data: ReportData,
        output_path: Optional[str]
    ) -> str:
        """Generate HTML report."""
        # Create HTML template if not exists
        html_template_content = self._get_html_template(template)
        
        # Render template
        jinja_template = Template(html_template_content)
        html_content = jinja_template.render(
            template=template,
            report=report_data,
            generated_at=datetime.utcnow(),
            sections=self._get_ordered_sections(template, report_data)
        )
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"HTML report saved to: {output_path}")
        
        return html_content
    
    async def _generate_pdf_report(
        self,
        template: ReportTemplate,
        report_data: ReportData,
        output_path: Optional[str]
    ) -> str:
        """Generate PDF report."""
        # First generate HTML
        html_content = await self._generate_html_report(template, report_data, None)
        
        # Convert to PDF
        try:
            # Get CSS for PDF styling
            css_content = self._get_pdf_css(template)
            
            # Generate PDF
            html_doc = HTML(string=html_content)
            css_doc = CSS(string=css_content) if css_content else None
            
            if output_path:
                if css_doc:
                    html_doc.write_pdf(output_path, stylesheets=[css_doc])
                else:
                    html_doc.write_pdf(output_path)
                logger.info(f"PDF report saved to: {output_path}")
                return output_path
            else:
                # Return PDF as bytes (base64 encoded for transport)
                import base64
                if css_doc:
                    pdf_bytes = html_doc.write_pdf(stylesheets=[css_doc])
                else:
                    pdf_bytes = html_doc.write_pdf()
                return base64.b64encode(pdf_bytes).decode('utf-8')
                
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            # Fallback to HTML
            if output_path:
                html_path = output_path.replace('.pdf', '.html')
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return html_path
            return html_content
    
    async def _generate_docx_report(
        self,
        template: ReportTemplate,
        report_data: ReportData,
        output_path: Optional[str]
    ) -> str:
        """Generate DOCX report."""
        doc = Document()
        
        # Add title
        title = doc.add_heading(report_data.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if report_data.subtitle:
            subtitle = doc.add_heading(report_data.subtitle, 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {report_data.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Generated by: {report_data.generated_by}")
        
        if report_data.company_name:
            doc.add_paragraph(f"Company: {report_data.company_name}")
        
        doc.add_page_break()
        
        # Add table of contents placeholder
        if template.include_toc:
            doc.add_heading("Table of Contents", 1)
            doc.add_paragraph("(Table of contents would be generated here)")
            doc.add_page_break()
        
        # Add sections
        sections = self._get_ordered_sections(template, report_data)
        for section in sections:
            if section.page_break_before:
                doc.add_page_break()
            
            # Add section heading
            doc.add_heading(section.title, section.heading_level)
            
            # Add section content
            if section.content:
                # Split content into paragraphs
                paragraphs = section.content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Add subsections
            for subsection in section.subsections:
                doc.add_heading(subsection.title, subsection.heading_level)
                if subsection.content:
                    paragraphs = subsection.content.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
        
        # Add appendices if enabled
        if template.include_appendices and report_data.appendices:
            doc.add_page_break()
            doc.add_heading("Appendices", 1)
            
            for appendix_name, appendix_content in report_data.appendices.items():
                doc.add_heading(appendix_name, 2)
                if isinstance(appendix_content, str):
                    doc.add_paragraph(appendix_content)
                elif isinstance(appendix_content, dict):
                    doc.add_paragraph(json.dumps(appendix_content, indent=2))
        
        # Save document
        if output_path:
            doc.save(output_path)
            logger.info(f"DOCX report saved to: {output_path}")
            return output_path
        else:
            # Save to temporary file and return path
            import tempfile
            temp_path = tempfile.mktemp(suffix='.docx')
            doc.save(temp_path)
            return temp_path
    
    async def _generate_markdown_report(
        self,
        template: ReportTemplate,
        report_data: ReportData,
        output_path: Optional[str]
    ) -> str:
        """Generate Markdown report."""
        md_content = []
        
        # Title
        md_content.append(f"# {report_data.title}")
        if report_data.subtitle:
            md_content.append(f"## {report_data.subtitle}")
        
        md_content.append("")
        
        # Metadata
        md_content.append("---")
        md_content.append(f"**Generated:** {report_data.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        md_content.append(f"**Generated by:** {report_data.generated_by}")
        if report_data.company_name:
            md_content.append(f"**Company:** {report_data.company_name}")
        if report_data.ticker_symbol:
            md_content.append(f"**Ticker:** {report_data.ticker_symbol}")
        md_content.append("---")
        md_content.append("")
        
        # Table of contents
        if template.include_toc:
            md_content.append("## Table of Contents")
            md_content.append("")
            sections = self._get_ordered_sections(template, report_data)
            for i, section in enumerate(sections, 1):
                md_content.append(f"{i}. [{section.title}](#{section.section_id})")
            md_content.append("")
        
        # Sections
        sections = self._get_ordered_sections(template, report_data)
        for section in sections:
            # Section heading
            heading_prefix = "#" * (section.heading_level + 1)
            md_content.append(f"{heading_prefix} {section.title} {{#{section.section_id}}}")
            md_content.append("")
            
            # Section content
            if section.content:
                md_content.append(section.content)
                md_content.append("")
            
            # Subsections
            for subsection in section.subsections:
                sub_heading_prefix = "#" * (subsection.heading_level + 1)
                md_content.append(f"{sub_heading_prefix} {subsection.title}")
                md_content.append("")
                if subsection.content:
                    md_content.append(subsection.content)
                    md_content.append("")
        
        # Appendices
        if template.include_appendices and report_data.appendices:
            md_content.append("## Appendices")
            md_content.append("")
            
            for appendix_name, appendix_content in report_data.appendices.items():
                md_content.append(f"### {appendix_name}")
                md_content.append("")
                if isinstance(appendix_content, str):
                    md_content.append(appendix_content)
                elif isinstance(appendix_content, dict):
                    md_content.append("```json")
                    md_content.append(json.dumps(appendix_content, indent=2))
                    md_content.append("```")
                md_content.append("")
        
        markdown_text = "\n".join(md_content)
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_text)
            logger.info(f"Markdown report saved to: {output_path}")
        
        return markdown_text
    
    async def _generate_json_report(
        self,
        template: ReportTemplate,
        report_data: ReportData,
        output_path: Optional[str]
    ) -> str:
        """Generate JSON report."""
        # Convert report data to JSON-serializable format
        json_data = {
            "template": {
                "template_id": template.template_id,
                "name": template.name,
                "type": template.template_type.value
            },
            "report": {
                "title": report_data.title,
                "subtitle": report_data.subtitle,
                "company_name": report_data.company_name,
                "ticker_symbol": report_data.ticker_symbol,
                "report_type": report_data.report_type,
                "generated_at": report_data.generated_at.isoformat(),
                "generated_by": report_data.generated_by,
                "report_id": report_data.report_id
            },
            "content": {
                "executive_summary": report_data.executive_summary,
                "key_findings": report_data.key_findings,
                "recommendations": report_data.recommendations,
                "risk_factors": report_data.risk_factors
            },
            "sections": {},
            "data": {
                "financial_metrics": report_data.financial_metrics,
                "market_data": report_data.market_data,
                "data_sources": report_data.data_sources,
                "data_quality_score": report_data.data_quality_score
            },
            "appendices": report_data.appendices,
            "metadata": {
                "word_count": report_data.calculate_total_word_count(),
                "section_count": len(report_data.sections)
            }
        }
        
        # Add sections
        for section_id, section in report_data.sections.items():
            json_data["sections"][section_id] = {
                "title": section.title,
                "content": section.content,
                "order": section.order,
                "heading_level": section.heading_level,
                "word_count": section.word_count,
                "data_sources": section.data_sources,
                "confidence_score": section.confidence_score,
                "subsections": [
                    {
                        "title": sub.title,
                        "content": sub.content,
                        "word_count": sub.word_count
                    }
                    for sub in section.subsections
                ]
            }
        
        json_text = json.dumps(json_data, indent=2, ensure_ascii=False)
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_text)
            logger.info(f"JSON report saved to: {output_path}")
        
        return json_text
    
    def _get_ordered_sections(
        self,
        template: ReportTemplate,
        report_data: ReportData
    ) -> List[ReportSection]:
        """Get sections in the order specified by the template."""
        ordered_sections = []
        
        for section_id in template.sections:
            section = report_data.get_section(section_id)
            if section:
                ordered_sections.append(section)
            else:
                # Create placeholder section if missing
                placeholder = ReportSection(
                    section_id=section_id,
                    title=section_id.replace('_', ' ').title(),
                    content=f"[{section_id} content not available]"
                )
                ordered_sections.append(placeholder)
        
        # Sort by order field
        ordered_sections.sort(key=lambda s: s.order)
        
        return ordered_sections
    
    def _get_html_template(self, template: ReportTemplate) -> str:
        """Get HTML template content."""
        if template.html_template and (self.templates_dir / template.html_template).exists():
            with open(self.templates_dir / template.html_template, 'r', encoding='utf-8') as f:
                return f.read()
        
        # Return default HTML template
        return self._get_default_html_template()
    
    def _get_default_html_template(self) -> str:
        """Get default HTML template."""
        return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ report.title }}</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }
        .header { text-align: center; margin-bottom: 40px; }
        .title { font-size: 2.5em; color: #2c3e50; margin-bottom: 10px; }
        .subtitle { font-size: 1.5em; color: #7f8c8d; margin-bottom: 20px; }
        .metadata { background: #f8f9fa; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .section { margin-bottom: 30px; }
        .section-title { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 5px; }
        .toc { background: #f8f9fa; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .toc ul { list-style-type: none; padding-left: 0; }
        .toc li { margin: 5px 0; }
        .toc a { text-decoration: none; color: #3498db; }
        .executive-summary { background: #e8f4fd; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .key-findings { background: #f0f8f0; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .recommendations { background: #fff3cd; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .risk-factors { background: #f8d7da; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .footer { margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd; text-align: center; color: #7f8c8d; }
    </style>
</head>
<body>
    <div class="header">
        <h1 class="title">{{ report.title }}</h1>
        {% if report.subtitle %}
        <h2 class="subtitle">{{ report.subtitle }}</h2>
        {% endif %}
    </div>
    
    <div class="metadata">
        <p><strong>Generated:</strong> {{ report.generated_at.strftime('%Y-%m-%d %H:%M:%S') }}</p>
        <p><strong>Generated by:</strong> {{ report.generated_by }}</p>
        {% if report.company_name %}
        <p><strong>Company:</strong> {{ report.company_name }}</p>
        {% endif %}
        {% if report.ticker_symbol %}
        <p><strong>Ticker:</strong> {{ report.ticker_symbol }}</p>
        {% endif %}
    </div>
    
    {% if template.include_toc %}
    <div class="toc">
        <h2>Table of Contents</h2>
        <ul>
        {% for section in sections %}
            <li><a href="#{{ section.section_id }}">{{ section.title }}</a></li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
    
    {% if report.executive_summary %}
    <div class="executive-summary">
        <h2>Executive Summary</h2>
        <p>{{ report.executive_summary }}</p>
    </div>
    {% endif %}
    
    {% for section in sections %}
    <div class="section" id="{{ section.section_id }}">
        <h{{ section.heading_level + 1 }} class="section-title">{{ section.title }}</h{{ section.heading_level + 1 }}>
        {% if section.content %}
        <div class="section-content">
            {{ section.content | replace('\n', '<br>') | safe }}
        </div>
        {% endif %}
        
        {% for subsection in section.subsections %}
        <div class="subsection">
            <h{{ subsection.heading_level + 1 }}>{{ subsection.title }}</h{{ subsection.heading_level + 1 }}>
            {% if subsection.content %}
            <div class="subsection-content">
                {{ subsection.content | replace('\n', '<br>') | safe }}
            </div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endfor %}
    
    {% if report.key_findings %}
    <div class="key-findings">
        <h2>Key Findings</h2>
        <ul>
        {% for finding in report.key_findings %}
            <li>{{ finding }}</li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
    
    {% if report.recommendations %}
    <div class="recommendations">
        <h2>Recommendations</h2>
        <ul>
        {% for recommendation in report.recommendations %}
            <li>{{ recommendation }}</li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
    
    {% if report.risk_factors %}
    <div class="risk-factors">
        <h2>Risk Factors</h2>
        <ul>
        {% for risk in report.risk_factors %}
            <li>{{ risk }}</li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
    
    {% if template.include_appendices and report.appendices %}
    <div class="appendices">
        <h2>Appendices</h2>
        {% for appendix_name, appendix_content in report.appendices.items() %}
        <div class="appendix">
            <h3>{{ appendix_name }}</h3>
            {% if appendix_content is string %}
            <p>{{ appendix_content }}</p>
            {% else %}
            <pre>{{ appendix_content | tojson(indent=2) }}</pre>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    <div class="footer">
        <p>This report was generated automatically by the AI Investment Research System.</p>
        <p>Data sources: {{ report.data_sources | join(', ') if report.data_sources else 'Various' }}</p>
    </div>
</body>
</html>
        """.strip()
    
    def _get_pdf_css(self, template: ReportTemplate) -> str:
        """Get CSS for PDF styling."""
        return """
        @page {
            size: A4;
            margin: 2cm;
        }
        
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.5;
        }
        
        .title {
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin-bottom: 20pt;
        }
        
        .subtitle {
            font-size: 18pt;
            text-align: center;
            margin-bottom: 30pt;
        }
        
        .section-title {
            font-size: 16pt;
            font-weight: bold;
            margin-top: 20pt;
            margin-bottom: 10pt;
            page-break-after: avoid;
        }
        
        .page-break {
            page-break-before: always;
        }
        
        .toc {
            page-break-after: always;
        }
        
        .executive-summary {
            background: #f0f8ff;
            padding: 15pt;
            border: 1pt solid #ddd;
            margin: 15pt 0;
        }
        
        .metadata {
            background: #f8f8f8;
            padding: 10pt;
            border: 1pt solid #ccc;
            margin-bottom: 20pt;
        }
        """


# Global report template engine instance
report_template_engine = ReportTemplateEngine()